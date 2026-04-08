"""Carga de documentos desde archivos PDF, TXT y DOCX."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class Document:
    """Representación interna de un documento cargado."""

    text: str
    metadata: dict = field(default_factory=dict)


def load_txt(path: Path) -> Document:
    text = path.read_text(encoding="utf-8")
    return Document(text=text, metadata={"source": path.name, "type": "txt"})


def load_pdf(path: Path) -> Document:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    page_map: list[dict] = []
    offset = 0
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        start = offset
        parts.append(page_text)
        offset += len(page_text) + 1  # +1 por el \n del join
        page_map.append({"page": i + 1, "start": start, "end": offset - 1})

    text = "\n".join(parts)
    return Document(
        text=text,
        metadata={
            "source": path.name,
            "type": "pdf",
            "total_pages": len(reader.pages),
            "page_map": page_map,
        },
    )


def load_docx(path: Path) -> Document:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return Document(
        text=text,
        metadata={"source": path.name, "type": "docx", "paragraphs": len(paragraphs)},
    )


_LOADERS: dict[str, callable] = {
    ".txt": load_txt,
    ".pdf": load_pdf,
    ".docx": load_docx,
}

SUPPORTED_EXTENSIONS = set(_LOADERS.keys())


def load_document(path: Path) -> Document:
    """Carga un documento según su extensión.

    Raises:
        ValueError: Si la extensión no es soportada.
    """
    ext = path.suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(
            f"Extensión '{ext}' no soportada. Usa: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return loader(path)
